from pathlib import Path
from docx import Document

source = Path(r"C:\Users\jap\Downloads\Research paper group 27 (2).docx")
document = Document(source)
lines = []

for index, paragraph in enumerate(document.paragraphs, start=1):
    text = paragraph.text.strip()
    if text:
        lines.append(f"P{index:04d}\t[{paragraph.style.name}]\t{text}")

for table_index, table in enumerate(document.tables, start=1):
    lines.append(f"TABLE {table_index}")
    for row_index, row in enumerate(table.rows, start=1):
        values = [cell.text.replace("\n", " | ").strip() for cell in row.cells]
        lines.append(f"T{table_index:02d}R{row_index:03d}\t" + "\t".join(values))

output = Path(r"C:\Users\jap\Documents\mid_eval\.doc_review\paper_text.txt")
output.write_text("\n".join(lines), encoding="utf-8")
print(f"paragraphs={len(document.paragraphs)} tables={len(document.tables)} lines={len(lines)} output={output}")
